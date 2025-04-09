import sys
from docx import Document

#
# Word ファイルのテンプレート(契約書)に情報を書き込むスクリプト例
#

# コマンドライン引数
args = sys.argv

# 被雇用者
seller = args[1]

# Word ドキュメントのオブジェクトを生成する
filename = "23665_koyou_keiyaku.docx"
doc = Document(filename)

# 先頭のテーブル
tbl = doc.tables[0]

# 被雇用者氏名を入れる
tbl.rows[0].cells[1].text = seller
tbl.rows[0].cells[1]._add_paragraph()

# 生年月日
tbl.rows[0].cells[3].text = '1900年1月1日'

# 現住所
tbl.rows[1].cells[1].text = '〒XXX-YYYY 東京都千代田区A-B-C'
tbl.rows[1].cells[1].add_paragraph('あいうえお101')
tbl.rows[1].cells[1].add_paragraph('TEL aaa-bbbb-cccc')

# 新しく保存する
doc.save('new_keiyakusho.docx')